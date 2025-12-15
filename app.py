import streamlit as st
import google.generativeai as genai
from pypdf import PdfReader
from docx import Document
from io import BytesIO

# --- 1. CONFIGURATION DE LA PAGE ---
st.set_page_config(
    page_title="Assistant IA - Appels d'Offres BTP",
    page_icon="🏗️",
    layout="centered"
)

# --- 2. SÉCURITÉ & LOGIN (Le "Mur") ---
def check_password():
    """Retourne True si l'utilisateur est connecté."""
    if st.session_state.get('password_correct', False):
        return True

    # Interface de connexion
    st.markdown("## 🔒 Accès Réservé")
    st.markdown("Veuillez entrer votre clé d'accès client pour utiliser l'outil.")
    
    password = st.text_input("Code d'accès", type="password")
    
    if password:
        # Vérification via les secrets du serveur
        if password == st.secrets["ADMIN_PASSWORD"]:
            st.session_state["password_correct"] = True
            st.rerun() # Recharge la page pour afficher l'app
        else:
            st.error("❌ Code d'accès incorrect.")
    return False

if not check_password():
    st.stop() # Arrête tout si pas connecté

# --- 3. FONCTIONS MÉTIER (Le Moteur) ---

def get_api_key():
    """Récupère la clé API Google sécurisée."""
    try:
        return st.secrets["GOOGLE_API_KEY"]
    except:
        st.error("Erreur de configuration : Clé API manquante dans les secrets.")
        st.stop()

def extract_text_from_pdf(uploaded_file):
    """Extrait le texte brut du PDF."""
    try:
        pdf_reader = PdfReader(uploaded_file)
        text = ""
        for page in pdf_reader.pages:
            content = page.extract_text()
            if content:
                text += content
        return text
    except Exception as e:
        st.error(f"Erreur de lecture PDF : {e}")
        return None

def generate_analysis_gemini(text_content):
    api_key = get_api_key()
    genai.configure(api_key=api_key)
    
    # On utilise toujours Flash pour la vitesse et le contexte long
    model = genai.GenerativeModel('gemini-1.5-flash')
    
    prompt = f"""
    Tu es un Directeur Technique chevronné dans le BTP. Tu analyses un CCTP pour préparer un Mémoire Technique gagnant.
    
    DOCUMENT À ANALYSER :
    {text_content} 
    
    TA MISSION : 
    Ne fais pas de résumé général. Extrais uniquement les points critiques qui impactent le chiffrage et l'organisation.
    
    FORMAT DE RÉPONSE OBLIGATOIRE (Respecte cette structure pour le Word) :

    # 1. FICHE D'IDENTITÉ DU CHANTIER
    * **Nature des travaux :** (Ex: Rénovation thermique, Gros œuvre...)
    * **Contraintes de site majeures :** (Accès, stockage, horaires, site occupé ?)
    * **Délais & Planning :** (Dates clés ou durées mentionnées)

    # 2. POINTS DE VIGILANCE & PIÈGES (Crucial)
    * *Liste ici les éléments qui coûtent cher ou qu'on risque d'oublier.*
    * (Ex: Marques imposées, normes spécifiques DTU citées, performances acoustiques/thermiques exactes à atteindre).
    
    # 3. MOYENS TECHNIQUES SPÉCIFIQUES REQUIS
    * Ne mets pas "Outillage standard".
    * Cite les engins ou matériels lourds obligatoires selon le texte (Ex: Échafaudage classe 4, Grue, Cantonnement spécifique).

    # 4. ÉBAUCHE DU MÉMOIRE TECHNIQUE (Partie Rédigée)
    *Rédige un paragraphe argumentaire professionnel pour rassurer le client sur ces 2 points :*
    * **Notre méthodologie pour ce chantier :** (Adapte le texte aux contraintes identifiées plus haut).
    * **Gestion de la sécurité et environnement :** (Cite les obligations du CCTP : tri des déchets, nuisances sonores).

    TON : Direct, Technique, "Pro". Pas de blabla.
    """
    
    try:
        response = model.generate_content(prompt)
        return response.text
    except Exception as e:
        st.error(f"Erreur IA : {e}")
        return None

def create_word_doc(text_ia):
    """Transforme le texte de l'IA en fichier .docx téléchargeable."""
    doc = Document()
    doc.add_heading('Mémoire Technique - Ébauche IA', 0)
    doc.add_paragraph("Document généré automatiquement. À relire et compléter.")
    doc.add_paragraph("-" * 50)
    
    # On ajoute le contenu généré
    doc.add_paragraph(text_ia)
    
    # Sauvegarde en mémoire tampon (RAM)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 4. INTERFACE UTILISATEUR (SaaS) ---

# Gestion du quota (Anti-Abus)
if 'usage_count' not in st.session_state:
    st.session_state['usage_count'] = 0
QUOTA_MAX = 5 # Limite par session

st.title("🏗️ Générateur de Mémoire Technique")
st.caption("Solution IA pour artisans du bâtiment - Version Bêta")

st.info(f"💡 Crédits restants pour cette session : {QUOTA_MAX - st.session_state['usage_count']}")

uploaded_file = st.file_uploader("Déposez votre CCTP (PDF uniquement)", type="pdf")

if uploaded_file is not None:
    # Bouton d'action
    if st.button("🚀 Analyser et Générer le Word"):
        
        # 1. Vérif Quota
        if st.session_state['usage_count'] >= QUOTA_MAX:
            st.error("⚠️ Limite d'utilisation atteinte pour cette session.")
        else:
            st.session_state['usage_count'] += 1
            
            with st.spinner('Lecture du PDF et rédaction en cours (env. 20 secondes)...'):
                # 2. Extraction
                raw_text = extract_text_from_pdf(uploaded_file)
                
                if raw_text and len(raw_text) > 100:
                    # 3. Génération IA
                    analysis_result = generate_analysis_gemini(raw_text)
                    
                    if analysis_result:
                        st.success("Analyse terminée !")
                        
                        # 4. Affichage Aperçu
                        with st.expander("👁️ Voir l'aperçu du texte"):
                            st.markdown(analysis_result)
                        
                        # 5. Création Word
                        docx = create_word_doc(analysis_result)
                        
                        st.download_button(
                            label="📥 Télécharger le Mémoire (.docx)",
                            data=docx,
                            file_name="Memoire_Technique_Genere.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            type="primary" # Bouton mis en évidence
                        )
                else:
                    st.warning("Le PDF semble vide ou illisible (c'est peut-être une image scannée ?).")

st.markdown("---")
st.caption("© 2025 - Outil développé par [Votre Société] - Tous droits réservés.")